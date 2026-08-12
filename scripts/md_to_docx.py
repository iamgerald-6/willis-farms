#!/usr/bin/env python3
"""Convert platform markdown docs to a single Word (.docx) file."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUTPUT = DOCS / "Wills_Farms_Platform_Documentation.docx"
COLLEAGUE_OUTPUT = DOCS / "Colleague_Task_Brief.docx"

INPUT_FILES = [
    DOCS / "PLATFORM_AUDIT_AND_ROADMAP.md",
    DOCS / "SYSTEM_DEFINITIONS_SPEC.md",
]

COLLEAGUE_INPUT = DOCS / "COLLEAGUE_TASK_BRIEF.md"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str) -> None:
    """Parse inline **bold**, `code`, and [text](url)."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`\\[]+)"
    )
    for part in pattern.findall(text):
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.italic = True
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        else:
            paragraph.add_run(part)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def is_separator_row(line: str) -> bool:
    s = line.strip()
    if not is_table_row(s):
        return False
    cells = [c.strip() for c in s.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", c or "---") for c in cells)


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        if not is_separator_row(lines[i]):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            if ri == 0:
                run = p.add_run(text)
                run.bold = True
                set_cell_shading(cell, "E8E8E8")
            else:
                add_formatted_runs(p, text)
    doc.add_paragraph()


def convert_markdown(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("_" * 72)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped[level:].strip()
            level = min(max(level, 1), 4)
            doc.add_heading(title, level=level)
            i += 1
            continue

        if is_table_row(stripped):
            rows, i = parse_table_rows(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, text)
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, m.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1


def build_document_from_files(
    title: str,
    subtitle: str,
    input_files: list[Path],
) -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, 0)

    sub = doc.add_paragraph()
    sub.add_run(subtitle).italic = True
    doc.add_paragraph("Generated from repository documentation — August 2026.")
    doc.add_paragraph()

    for idx, path in enumerate(input_files):
        if not path.exists():
            print(f"Warning: missing {path}", file=sys.stderr)
            continue
        if idx > 0:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        doc.add_heading(path.stem.replace("_", " "), level=1)
        md_text = path.read_text(encoding="utf-8")
        # Skip duplicate top-level title from md if present
        md_lines = md_text.splitlines()
        if md_lines and md_lines[0].startswith("# "):
            md_text = "\n".join(md_lines[1:])
        convert_markdown(doc, md_text)

    return doc


def build_document() -> Document:
    return build_document_from_files(
        "Wills Farms Platform Documentation",
        "Infrastructure Audit, Roadmap & System Definitions",
        INPUT_FILES,
    )


def build_colleague_document() -> Document:
    return build_document_from_files(
        "Colleague Task Brief",
        "UI & Form Automation — four scoped tasks",
        [COLLEAGUE_INPUT],
    )


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")

    if COLLEAGUE_INPUT.exists():
        colleague_doc = build_colleague_document()
        colleague_doc.save(COLLEAGUE_OUTPUT)
        print(f"Created: {COLLEAGUE_OUTPUT}")


if __name__ == "__main__":
    main()
